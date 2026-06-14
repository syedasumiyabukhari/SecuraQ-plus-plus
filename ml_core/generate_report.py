"""
Generate SecuraQpp v2 — Full Methodology & Experimental Report (Word .docx)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from pathlib import Path

ROOT   = Path(__file__).resolve().parent
CHARTS = ROOT / 'results/defense'
METS   = ROOT / 'results/metrics'
OUT    = ROOT / 'results/SecuraQpp_v2_Methodology_Report.docx'

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_img(path, width=6.0, caption=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size  = Pt(9)
            cp.runs[0].font.italic = True
            cp.runs[0].font.color.rgb = RGBColor(0x55,0x55,0x55)

def table_row(tbl, cells, bold=False, shade=None):
    row = tbl.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(txt)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9.5)
            if bold: run.font.bold = True
        if shade:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  shade)
            tcPr.append(shd)
    return row

def hdr_row(tbl, cols):
    table_row(tbl, cols, bold=True, shade='1A5376')
    for cell in tbl.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SecuraQpp v2')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1A,0x53,0x76)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Quantum-Enhanced Source Code Vulnerability Detection System')
r2.font.size   = Pt(16)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0x44,0x44,0x44)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Methodology, Normalization Steps & Experimental Report')
r3.font.size  = Pt(13)
r3.font.bold  = True

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('Final Year Project — Department of Computer Science\nMay 2026')
r4.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. Abstract')
body(
    'SecuraQpp v2 is a full-stack, quantum-classical hybrid system for automated software '
    'vulnerability detection. The pipeline processes C/C++ source code through eight sequential '
    'stages: text preprocessing, multi-view graph construction, graph attention network (GAT) '
    'encoding, classical compression, quantum-aware feature alignment (QAFA), variational quantum '
    'circuit (VQC) simulation, hybrid feature fusion, and final classification. Three vulnerability '
    'classes from the NIST Juliet Test Suite are targeted: Buffer Overflow (BO), Format String (FS), '
    'and Use-After-Free (UAF). Final test accuracies reached 90.5% (BO), 91.8% (UAF), and 82.0% '
    '(FS), all significantly outperforming the QEGVD paper baseline of 61.19%–78.6%. The system '
    'is deployed as a production web application with role-based access, two-factor authentication, '
    'real-time graph visualization, and AI-powered patch suggestions via the Claude API.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. DATASET
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. Dataset Description')

heading('2.1 Source', level=2)
body(
    'The NIST Juliet Test Suite (C/C++ version) is the industry-standard benchmark for '
    'vulnerability detection research. It contains thousands of synthetic C/C++ code pairs — '
    'each pair consists of one vulnerable function and one patched (safe) version — covering '
    'over 100 Common Weakness Enumeration (CWE) categories.'
)

heading('2.2 Selected Vulnerability Types', level=2)
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl, ['Vulnerability', 'CWE', 'Total Samples', 'Balance'])
for row in [
    ('Buffer Overflow (BO)',   'CWE-121/122', '11,417', '50% vuln / 50% safe'),
    ('Format String (FS)',     'CWE-134',     '5,390',  '50% vuln / 50% safe'),
    ('Use-After-Free (UAF)',   'CWE-416',     '891',    '50% vuln / 50% safe'),
]:
    table_row(tbl, row)
doc.add_paragraph()

heading('2.3 Preprocessing & Splits', level=2)
body('Each dataset undergoes the following before splitting:')
bullet('Variable names replaced with VAR_N tokens (prevents name-based leakage)')
bullet('Type names replaced with TYPE_N tokens')
bullet('Whitespace collapsed, duplicate code removed via MD5 hashing')
bullet('15-point data leakage audit run — aborts if CRITICAL/HIGH violation detected')
body('Stratified split applied: 70% Train / 15% Validation / 15% Test')

tbl2 = doc.add_table(rows=1, cols=5)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl2, ['Dataset', 'Train', 'Val', 'Test', 'Total'])
for row in [
    ('BO',  '7,991', '1,713', '1,713', '11,417'),
    ('FS',  '3,772', '809',   '809',  '5,390'),
    ('UAF', '623',   '134',   '134',  '891'),
]:
    table_row(tbl2, row)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. Methodology — Pipeline Stages')

body(
    'The QEGVD (Quantum-Enhanced Graph Vulnerability Detection) pipeline processes each code '
    'sample through 8 sequential stages. The architecture is summarized below:'
)

# Pipeline overview table
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl3, ['Stage', 'Name', 'Output Dimension'])
stages = [
    ('Stage 1', 'Data Preprocessing & Splitting',          'Cleaned CSVs, 70/15/15 splits'),
    ('Stage 2', 'Multi-View Graph Construction',           '8 graphs per sample, 64-dim node features'),
    ('Stage 3', 'Graph Attention Network (GAT)',           '256-dim graph embedding'),
    ('Stage 4', 'Classical Encoder (Compression)',         '256-dim vulnerability signature'),
    ('Stage 5', 'QAFA — Quantum-Aware Feature Alignment',  '8×8-dim = 64 angle-encoded features'),
    ('Stage 6', 'Variational Quantum Circuit (VQC)',       '4–40-dim quantum feature vector (per pipeline)'),
    ('Stage 7', 'Hybrid Feature Fusion',                   '260-dim (UAF) / 288-dim (BO) / 311-dim (FS)'),
    ('Stage 8', 'MLP Final Classifier',                    'Binary label + confidence score'),
]
for s in stages:
    table_row(tbl3, s)
doc.add_paragraph()

# ── Stage 1 ───────────────────────────────────────────────────────────────────
heading('3.1 Stage 1 — Data Preprocessing', level=2)
body(
    'The raw Juliet CSV files are loaded and passed through a rigorous cleaning pipeline '
    'before any model sees them.'
)
bullet('Whitespace Normalization: tabs → single space; 2+ spaces → single space; 3+ newlines → 2 newlines; strip leading/trailing whitespace')
bullet('MD5 Near-Duplicate Removal: each sample is lowercased, all whitespace collapsed to single space, then MD5-hashed. Duplicate hashes are dropped.')
bullet('15-Point Leakage Audit: checks for label-in-code, identifier overlap between splits, and other critical data contamination signals.')
bullet('Stratified Split: sklearn StratifiedKFold ensures class balance is maintained in all three splits.')

# ── Stage 2 ───────────────────────────────────────────────────────────────────
heading('3.2 Stage 2 — Multi-View Graph Construction', level=2)
body(
    'Each code sample is converted into 8 distinct program analysis graphs using pure Python '
    'static analysis (no external parser required — fully offline and portable).'
)
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl4, ['Graph', 'Full Name', 'What It Captures'])
for row in [
    ('AST', 'Abstract Syntax Tree',       'Syntactic code structure'),
    ('CFG', 'Control Flow Graph',         'Execution paths (if/else/loops)'),
    ('DFG', 'Data Flow Graph',            'Variable definitions and uses'),
    ('PDG', 'Program Dependence Graph',   'CFG + DFG unified'),
    ('TPG', 'Taint Propagation Graph',    'User input → dangerous sink flows'),
    ('MAG', 'Memory Access Graph',        'malloc/free/access relationships'),
    ('CG',  'Call Graph',                 'Caller → callee relationships'),
    ('FSG', 'Format String Graph',        'Format specifier → argument binding and taint paths'),
]:
    table_row(tbl4, row)
doc.add_paragraph()
body('Each node is assigned a 64-dimensional feature vector encoding token type, structural metrics, and vulnerability-specific signals.')

# ── Stage 3 ───────────────────────────────────────────────────────────────────
heading('3.3 Stage 3 — Graph Attention Network (GAT)', level=2)
body('Architecture: per-view encoder + attention-weighted fusion.')
bullet('Per view: 4× GATConv layers (128-dim, 8 attention heads each) → mean+max readout → per-view embedding')
bullet('Fusion: learned attention weights over all 8 views → 256-dim fused graph embedding (FUSED_DIM=256)')
bullet('Classification head: Linear(256→64→1) with Sigmoid')
bullet('Loss: Focal Loss + Supervised Contrastive Loss (SupCon), combined with α=0.7/0.3 weighting')
body('Per-dataset parameters:')
tbl5 = doc.add_table(rows=1, cols=4)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl5, ['Parameter', 'BO', 'FS', 'UAF'])
for row in [
    ('Dropout',         '0.30',  '0.20',  '0.30'),
    ('DropEdge',        '15%',   '10%',   '15%'),
    ('SupCon Temperature', '0.07', '0.10', '0.10'),
    ('Label Smoothing', '0.05',  '0.02',  '0.05'),
]:
    table_row(tbl5, row)
doc.add_paragraph()

# ── Stage 4 ───────────────────────────────────────────────────────────────────
heading('3.4 Stage 4 — Classical Encoder', level=2)
body('Compresses the 256-dim GAT embedding into a compact vulnerability signature.')
bullet('Architecture: FC(256→128) → ReLU → BN → Dropout → FC(128→256) → ReLU → BN (INPUT_DIM=256)')
bullet('Input is L2-normalized before the first FC layer')
bullet('Output: 256-dim vulnerability signature — this full vector feeds into Stage 7 fusion')
bullet('QAFA (Stage 5) takes this 256-dim output and selects top 16 features for the VQC only')
tbl6 = doc.add_table(rows=1, cols=4)
tbl6.style = 'Table Grid'
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl6, ['Pipeline', 'Input (GAT)', 'Architecture', 'Output'])
for row in [
    ('BO',  '256-dim', 'FC(256→128)→BN→FC(128→256)→BN', '256-dim'),
    ('FS',  '256-dim', 'FC(256→128)→BN→FC(128→256)→BN', '256-dim'),
    ('UAF', '256-dim', 'FC(256→128)→BN→FC(128→256)→BN', '256-dim'),
]:
    table_row(tbl6, row)
doc.add_paragraph()
body('All pipelines output 256-dim. The 256-dim embedding passes directly to Stage 7 fusion. QAFA selects a 16-feature subset for the quantum circuit only.')

# ── Stage 5 ───────────────────────────────────────────────────────────────────
heading('3.5 Stage 5 — QAFA (Quantum-Aware Feature Alignment)', level=2)
body(
    'QAFA selects the most vulnerability-relevant features from the compressed vector '
    'before passing them to the quantum circuit. This is critical: quantum circuits have '
    'limited expressibility per qubit, so only the most informative features should enter.'
)
bullet('Composite scoring: S_i = α·MI(i) + β·SHAP(i) + γ·Centrality(i)  [all normalized to 0–1]')
bullet('MI: Mutual Information with the label (sklearn, 5 nearest neighbours)')
bullet('SHAP: Mean absolute SHAP value from a GBM model')
bullet('Centrality: Weighted average of degree, betweenness, and closeness centrality across 200 sampled graphs')
bullet('Top 16 features selected → split into 2 groups of 8 → rescaled to [–π, π] via tanh×π for angle encoding')
body('FS-specific extras:')
bullet('Histogram Overlap Filter: eliminates features where class distributions overlap >90%')
bullet('MMD (Maximum Mean Discrepancy) as 4th criterion (δ=0.25 weight)')
tbl7 = doc.add_table(rows=1, cols=5)
tbl7.style = 'Table Grid'
tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl7, ['Pipeline', 'α (MI)', 'β (SHAP)', 'γ (Centrality)', 'δ (MMD)'])
for row in [
    ('BO',  '0.30', '0.20', '0.50', '—'),
    ('FS',  '0.30', '0.25', '0.20', '0.25'),
    ('UAF', '0.40', '0.35', '0.25', '—'),
]:
    table_row(tbl7, row)
doc.add_paragraph()

# ── Stage 6 ───────────────────────────────────────────────────────────────────
heading('3.6 Stage 6 — Variational Quantum Circuit (VQC)', level=2)
body(
    'A real quantum circuit simulation is implemented using PennyLane (default.qubit backend). '
    'The circuit operates on 4 qubits and uses the data re-uploading technique to process '
    '64 input features through 8 encoding rounds.'
)
bullet('Circuit: H⁴ → RY/RZ(s1) → CNOT-ring → Var_A(×3) → RY/RZ(s2) → CNOT-ring → Var_B(×3) → ⟨Z⟩⁴')
bullet('4 qubits, 3 variational layers per block, 60 trainable parameters (standard VQC)')
bullet('Data re-uploading: 8 rounds × 8-dim input = 64 features encoded quantum-mechanically')
bullet('Output: 32-dim quantum feature vector (Pauli-Z expectation values concatenated across rounds)')
bullet('Learnable input_scale parameter (initialized to π) scales each feature individually')
bullet('Training: AdamW optimizer, Focal Loss, Cosine Annealing LR, early stopping (patience=15)')

# ── Stage 7+8 ─────────────────────────────────────────────────────────────────
heading('3.7 Stages 7+8 — Hybrid Fusion & Classification', level=2)
body(
    'The full 256-dim classical embedding from Stage 4 is concatenated with the VQC quantum '
    'output from Stage 6. Final fusion dimensions differ per pipeline:'
)
tbl_fus = doc.add_table(rows=1, cols=4)
tbl_fus.style = 'Table Grid'
tbl_fus.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_fus, ['Pipeline', 'Classical (S4)', 'Quantum (S6)', 'Total Fused Dim'])
for row in [
    ('BO',  '256-dim', '32-dim',        '288-dim'),
    ('UAF', '256-dim', '4-dim',         '260-dim'),
    ('FS',  '256-dim', '40-dim + 15 meta', '311-dim'),
]:
    table_row(tbl_fus, row)
doc.add_paragraph()
body('Ensemble: XGBoost + GradientBoosting + RandomForest with blend weight optimization and threshold calibration on validation set.')
body('MLP classifier on fused vector — FS uses BatchNorm on both FC layers with higher dropout (0.35/0.25) due to its smaller training set.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. NORMALIZATION STEPS
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. Normalization Steps — Per Pipeline')

heading('4.1 Stage 1 — Text-Level Normalization', level=2)
body('Identical for all 3 pipelines.')
bullet('Whitespace: tabs → single space; 2+ spaces → single space; 3+ newlines → 2 newlines; strip edges')
bullet('MD5 Near-Duplicate: lowercase entire string, collapse all whitespace → single space, strip, compute MD5. Samples with identical hash → one dropped.')

heading('4.2 Stage 2 — Node Feature Normalization', level=2)
body('Applied inline during graph construction before GAT training:')
tbl8 = doc.add_table(rows=1, cols=4)
tbl8.style = 'Table Grid'
tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl8, ['Feature Dim', 'Description', 'Normalization', 'Pipelines'])
for row in [
    ('[35]',    'Token count',          'min(n_tokens/30, 1.0) — clamp to [0,1]',    'All'),
    ('[36]',    'Call density',         'paren_count / n_tokens — ratio [0,1]',       'All'),
    ('[37]',    'Array access density', 'bracket_count / n_tokens',                   'All'),
    ('[38]',    'Pointer density',      'star_count / n_tokens',                      'All'),
    ('[39]',    'Assignment density',   'equals_count / n_tokens',                    'All'),
    ('[40]',    'Statement position',   'stmt_index / total_statements',              'All'),
    ('[41–47]', 'Binary signals',       'Already 0/1 — no normalization',            'All'),
    ('[48]',    'Block depth',          'sum({,}) / 5.0, clipped to 1.0',            'All'),
    ('[71]',    'Sink danger score',    'Hand-scored: snprintf=0.3, sprintf=1.0',     'FS'),
    ('[74]',    'Format sink count',    'count / total_statements',                   'FS'),
    ('[75]',    'Taint-to-sink ratio',  'taint_count / max(sink_count, 1)',           'FS'),
    ('[80]',    'Largest numeric literal', 'max_literal / 1000',                     'BO'),
]:
    table_row(tbl8, row)
doc.add_paragraph()

body('FS-specific Group-Wise Node Feature Normalizer (fitted on training set only):')
tbl9 = doc.add_table(rows=1, cols=3)
tbl9.style = 'Table Grid'
tbl9.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl9, ['Dimensions', 'Scaler Used', 'Reason'])
for row in [
    ('[0:35] — one-hot node types',       'None',          'Already {0,1} — StandardScaler would distort binary probabilities'),
    ('[35:64] — structural/lexical',      'RobustScaler',  'Robust to outlier functions with very long loops or many calls'),
    ('[64:84] — sink & FS-specific',      'MinMaxScaler',  'Counts and binary signals that must stay non-negative'),
]:
    table_row(tbl9, row)
doc.add_paragraph()

body('FS-specific TPG Edge Weight Clipping: taint edges clipped from max=5.0 → max=2.0 to prevent gradient explosion in GAT message passing.')

heading('4.3 Stage 3 — GAT Internal Normalization', level=2)
bullet('LayerNorm after every GATConv layer — normalizes across feature dim per node (not batch), stabilizing attention across nodes with different degree counts')
bullet('LayerNorm on readout: after mean+max pooling concat → Linear → LayerNorm(out_dim×2)')
bullet('LayerNorm on fusion: after attention-weighted sum → Linear(256) → LayerNorm(256)')
bullet('L2 Normalization before SupCon Loss: F.normalize(features, p=2, dim=1) — mathematically required for temperature-scaled cosine similarity')

heading('4.4 Stage 4 — Classical Encoder Normalization', level=2)
bullet('L2 Normalization of input (first operation in forward pass): x = F.normalize(x, p=2, dim=1)')
bullet('BatchNorm1d after every FC layer: normalizes each feature dim across batch, prevents internal covariate shift')
bullet('All pipelines: FC(256→128)→BN→Dropout→FC(128→256)→BN  (INPUT_DIM=256 confirmed in stage4_classical_encoder.py:76)')

heading('4.5 Stage 5 — QAFA Score Normalization', level=2)
bullet('MI Normalization: mi_scores /= mi_scores.max()  →  [0,1]')
bullet('SHAP Normalization: |shap_values|.mean(axis=0) /= max  →  [0,1]')
bullet('Centrality — Two-Level: (1) normalize per sample by local max; (2) normalize accumulated average by global max')
bullet('Composite score automatically in [0,1] since all three inputs are [0,1] and weights sum to 1.0')
bullet('Angle Encoding: angles = tanh(features) × π  →  [–π, π]  (maps to full Bloch sphere range for RY/RZ gates)')
body('Note: tanh×π chosen over MinMax or StandardScaler because it preserves negative values and uses the full rotation range of the quantum gate, maximizing circuit expressibility.')

heading('4.6 Stage 6 — VQC Normalization', level=2)
bullet('Learnable input_scale = nn.Parameter(torch.ones(4) × π): scales each of 4 input features independently per qubit')
bullet('Initialized to π to match the angle encoding range; learned during training to tune per-feature Bloch sphere rotation')

heading('4.7 Stages 7+8 — Fusion & MLP Normalization', level=2)
bullet('BO/UAF MLP: BatchNorm1d on first FC layer only')
bullet('FS MLP: BatchNorm1d on both FC layers — extra regularization due to smaller training set (~3,700 samples)')
bullet('FS dropout: 0.35 (layer 1) and 0.25 (layer 2) — higher than BO/UAF (0.30/0.20) to prevent overfitting')

body('Complete normalization summary by pipeline:')
tbl10 = doc.add_table(rows=1, cols=4)
tbl10.style = 'Table Grid'
tbl10.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl10, ['Stage', 'Normalization', 'BO', 'FS/UAF'])
summary_rows = [
    ('S1', 'Whitespace collapse',                '✓', '✓'),
    ('S1', 'MD5 lowercase+whitespace dedup',     '✓', '✓'),
    ('S2', 'Token count /30 clamp',              '✓', '✓'),
    ('S2', 'Density features as ratios',         '✓', '✓'),
    ('S2', 'RobustScaler [35:64]',               '✗', 'FS only'),
    ('S2', 'MinMaxScaler [64:84]',               '✗', 'FS only'),
    ('S2', 'TPG edge weight clip (max=2.0)',      '✗', 'FS only'),
    ('S3', 'LayerNorm per GATConv',              '✓', '✓'),
    ('S3', 'LayerNorm on fused 256-dim',         '✓', '✓'),
    ('S3', 'L2-norm for SupCon Loss',            '✓', '✓'),
    ('S4', 'L2-norm on encoder input',           '✓', '✓'),
    ('S4', 'BatchNorm1d both FC layers',         '✓', '✓'),
    ('S5', 'MI/SHAP/Centrality max-norm →[0,1]','✓', '✓'),
    ('S5', 'Histogram overlap filter (>90%)',    '✗', 'FS only'),
    ('S5', 'MMD as 4th criterion',              '✗', 'FS only'),
    ('S5', 'tanh×π angle encoding',             '✓', '✓'),
    ('S6', 'Learnable input_scale (init=π)',     '✓', '✓'),
    ('S7+8','BatchNorm1d layer 1',              '✓', '✓'),
    ('S7+8','BatchNorm1d layer 2',              '✗', 'FS only'),
    ('S7+8','Higher dropout (0.35/0.25)',        '✗', 'FS only'),
]
for row in summary_rows:
    table_row(tbl10, row)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. EXPERIMENTAL SECTION
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. Experimental Section')

heading('5.1 Experimental Setup', level=2)
body('All experiments were run on a Windows 11 machine (Intel i7, 16GB RAM). The quantum circuit was simulated using PennyLane default.qubit (CPU-based statevector simulation). No real quantum hardware was used.')

tbl11 = doc.add_table(rows=1, cols=2)
tbl11.style = 'Table Grid'
tbl11.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl11, ['Component', 'Version / Detail'])
for row in [
    ('Python',       '3.10'),
    ('PyTorch',      '2.x'),
    ('PennyLane',    '0.44.1'),
    ('XGBoost',      '2.x'),
    ('scikit-learn', '1.x'),
    ('Optimizer',    'AdamW (lr=1e-3, wd=1e-4)'),
    ('LR Schedule',  'Cosine Annealing (T_max=100)'),
    ('Loss',         'Focal Loss (γ=2.0, α=0.25) + SupCon'),
    ('Early Stop',   'patience=15 on validation loss'),
    ('Batch Size',   '32 (GAT/VQC), 256 (Encoder)'),
]:
    table_row(tbl11, row)
doc.add_paragraph()

heading('5.2 Experiment 1 — GAT Baseline (Stage 3)', level=2)
body('Goal: Establish a graph-based baseline using only the 7-view GAT encoder without quantum enhancement.')
tbl12 = doc.add_table(rows=1, cols=5)
tbl12.style = 'Table Grid'
tbl12.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl12, ['Dataset', 'Accuracy', 'F1', 'AUC', 'MCC'])
for row in [
    ('BO',  '90.4%', '0.910', '0.971', '0.817'),
    ('FS',  '—',     '—',     '—',     '—'),
    ('UAF', '84.3%', '0.829', '0.927', '0.687'),
]:
    table_row(tbl12, row)
doc.add_paragraph()
body('Observation: GAT alone achieves strong results on BO and UAF. The multi-view graph representation captures structural vulnerability patterns effectively.')

heading('5.3 Experiment 2 — Classical Encoder Compression (Stage 4)', level=2)
body('Goal: Test whether compressing 512-dim GAT embeddings to 32–64-dim retains discriminative power.')
body('Result: Compression to 64-dim (BO/FS) and 32-dim (UAF) maintained near-identical accuracy to Stage 3, confirming that the vulnerability signal is concentrated in a low-dimensional subspace. BatchNorm and L2 normalization were essential — removing them caused ~3–4% accuracy drop.')

heading('5.4 Experiment 3 — QAFA Weight Allocation (8 Configurations)', level=2)
body(
    'Goal: Find the optimal weight allocation (γ, α, β) for the QAFA composite scoring formula '
    'S_i = α·MI + β·SHAP + γ·Centrality, and determine whether a 64-dim or 90-dim encoder '
    'produces better downstream results. 8 configurations were tested on the BO pipeline.'
)

tbl_qa = doc.add_table(rows=1, cols=6)
tbl_qa.style = 'Table Grid'
tbl_qa.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_qa, ['Exp', 'Centrality (γ)', 'MI (α)', 'SHAP (β)', 'Encoder', 'Hybrid Dim'])
for row in [
    ('v1', '0.50', '0.30', '0.20', '64-dim', '288'),
    ('v2', '0.60', '0.20', '0.20', '64-dim', '288'),
    ('v4', '0.65', '0.25', '0.10', '90-dim', '292'),
    ('v5', '0.70', '0.20', '0.10', '90-dim', '292'),
    ('v6', '0.60', '0.20', '0.20', '90-dim', '292'),
    ('v7', '0.50', '0.25', '0.25', '90-dim', '292'),
    ('v8', '0.50', '0.40', '0.10', '90-dim', '292'),
    ('v9', '0.50', '0.40', '0.10', '64-dim', '288'),
]:
    table_row(tbl_qa, row)
doc.add_paragraph()

body('Results across all 8 configurations:')
tbl_qr = doc.add_table(rows=1, cols=9)
tbl_qr.style = 'Table Grid'
tbl_qr.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_qr, ['Exp', 'Acc', 'F1', 'Precision', 'Recall', 'AUC', 'MCC', 'FPR', 'FNR'])
for row in [
    ('v1','0.9054','0.9117','0.8513','0.9812','0.9723','0.8205','0.1696','0.0188'),
    ('v2','0.9025','0.9096','0.8442','0.9859','0.9709','0.8166','0.1800','0.0141'),
    ('v4','0.9019','0.9077','0.8533','0.9695','0.9641','0.8115','0.1649','0.0305'),
    ('v5','0.9019','0.9081','0.8504','0.9742','0.9541','0.8125','0.1696','0.0258'),
    ('v6','0.9037','0.9094','0.8545','0.9718','0.9522','0.8151','0.1638','0.0282'),
    ('v7','0.9002','0.9068','0.8464','0.9765','0.9581','0.8100','0.1754','0.0235'),
    ('v8','0.9013','0.9042','0.8740','0.9366','0.9605','0.8048','0.1336','0.0634'),
    ('v9','0.9043','0.9103','0.8525','0.9765','0.9708','0.8173','0.1672','0.0235'),
]:
    table_row(tbl_qr, row)
doc.add_paragraph()

body('Key findings from the QAFA weight experiment:')
bullet('v1 is the best overall — highest F1 (91.17%), AUC (97.23%), MCC (0.82), lowest FNR (1.88%). For vulnerability detection, low FNR is critical: missing a real vulnerability is worse than a false alarm.')
bullet('64-dim encoder beats 90-dim — v1/v2/v9 (64-dim) consistently outperform v4–v8 (90-dim). The smaller, more compressed representation is more discriminative for the quantum circuit.')
bullet('Centrality sweet spot is 0.50 — increasing γ beyond 0.50 (v4 at 0.65, v5 at 0.70) degrades all metrics. Graph structure alone is insufficient; MI and SHAP are essential to balance it.')
bullet('v8 anomaly — v8 achieves the best precision (87.4%) and lowest FPR (13.36%) but the worst FNR (6.34%). It misses more real vulnerabilities — unacceptable for a security tool.')
bullet('v9 validates v1 architecture — v9 uses the same QAFA weights as v8 but with the 64-dim encoder. It recovers most of the performance lost in v8, confirming the 64-dim encoder as the key bottleneck.')
body('Final selected configuration: v1 — γ=0.50, α=0.30, β=0.20, 64-dim encoder.')

heading('5.5 Experiment 4 — FS Pipeline: 5 Experimental Runs', level=2)
body(
    'The FS pipeline was the hardest to improve — F1 was stuck at 65–72% across early runs '
    'due to VAR_N sanitization removing taint-tracking signals. Five full pipeline runs were '
    'executed with specific changes each time to systematically push performance.'
)

# Run 1
heading('Run 1 — Baseline (32 features, 4 rounds)', level=3)
body('Configuration: GAT retrained from scratch. QAFA: top-32 features, 4 rounds × 8 = 32Q features.')
tbl_r1 = doc.add_table(rows=1, cols=5)
tbl_r1.style = 'Table Grid'
tbl_r1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_r1, ['Stage', 'Accuracy', 'F1', 'AUC', 'MCC'])
for row in [
    ('Stage 3 (GAT)',     '0.6440', '0.6555', '0.6939', '0.2884'),
    ('Stage 4 (Encoder)', '0.6156', '0.7149', '0.6906', '0.3172'),
]:
    table_row(tbl_r1, row)
doc.add_paragraph()
body('Problem: GAT F1 only 65.5% — FS graph structure not discriminative enough with 32 features. Encoder F1 improved slightly to 71.5% but accuracy dropped.')

# Run 2
heading('Run 2 — Retrained GAT + 64 Features (8 rounds)', level=3)
body('Change: QAFA top-32 → top-64 features, 8 rounds. QAFA weights: centrality=0.25, MI=0.40, SHAP=0.35.')
tbl_r2 = doc.add_table(rows=1, cols=7)
tbl_r2.style = 'Table Grid'
tbl_r2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_r2, ['Stage', 'Accuracy', 'F1', 'AUC', 'MCC', 'FPR', 'FNR'])
for row in [
    ('Stage 3 (GAT)',     '0.6724', '0.7129', '0.7064', '0.3580', '0.4665', '0.1897'),
    ('Stage 4 (Encoder)', '0.6588', '0.7256', '0.7046', '0.3608', '0.5831', '0.1010'),
    ('Stage 6 (VQC)',     '0.6514', '0.7273', '0.6822', '0.3609', '0.6253', '0.0739'),
]:
    table_row(tbl_r2, row)
doc.add_paragraph()
body('Improvement: GAT F1 went from 65.5% → 71.3%. VQC F1 reached 72.7% but FPR was very high (62.5%) — the model was predicting too many false positives.')

# Run 3
heading('Run 3 — Full Pipeline with 64Q (Run 4)', level=3)
body('Change: Retrained hybrid MLP on top of Run 2 VQC output. Same 64-feature QAFA.')
tbl_r3 = doc.add_table(rows=1, cols=7)
tbl_r3.style = 'Table Grid'
tbl_r3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_r3, ['Stage', 'Accuracy', 'F1', 'AUC', 'MCC', 'FPR', 'FNR'])
for row in [
    ('Stage 8 (Hybrid Final)', '0.6712', '0.7246', '0.7160', '0.3694', '0.5211', '0.1379'),
]:
    table_row(tbl_r3, row)
doc.add_paragraph()
body('FPR improved from 62.5% → 52.1%. AUC improved to 71.6%. The hybrid MLP helped balance precision-recall but accuracy remained low.')

# Run 4
heading('Run 4 — 303-dim Hybrid Fusion (Run 5)', level=3)
body('Change: Hybrid fusion dim changed from 288 → 303-dim (different QAFA feature selection giving more quantum signal).')
tbl_r4 = doc.add_table(rows=1, cols=7)
tbl_r4.style = 'Table Grid'
tbl_r4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_r4, ['Stage', 'Accuracy', 'F1', 'AUC', 'MCC', 'FPR', 'FNR'])
for row in [
    ('Stage 8 (Hybrid Final)', '0.6761', '0.7277', '0.7131', '0.3784', '0.5112', '0.1379'),
]:
    table_row(tbl_r4, row)
doc.add_paragraph()
body('Best FS result from quantum pipeline runs — saved as intermediate model. FPR reduced to 51.1%, MCC improved to 0.378.')

# Run 5 — Fusion with supervisor improvements
heading('Run 5 — Full Fusion with Supervisor Improvements (Final)', level=3)
body('Changes: Enriched dataset features (supervisor), 64-feature VQC (supervisor), rule-based features, XGBoost ensemble + threshold calibration.')
tbl_r5 = doc.add_table(rows=1, cols=7)
tbl_r5.style = 'Table Grid'
tbl_r5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_r5, ['Stage', 'Accuracy', 'F1', 'AUC', 'MCC', 'FPR', 'FNR'])
for row in [
    ('Stage 8 (Final Fusion)', '0.8195', '0.8274', '0.8900', '0.6412', '0.2233', '0.1379'),
]:
    table_row(tbl_r5, row)
doc.add_paragraph()
body('Supervisor contributions that enabled this jump:')
bullet('Enriched dataset: code-level metrics (complexity, nesting depth, token ratios) added as supplementary features alongside graph embeddings.')
bullet('64-feature VQC: supervisor instructed to increase to 8-round re-uploading (64 features) to match BO and UAF configurations.')
bullet('Tokenizer features: format-string-specific patterns (printf/fprintf/snprintf, VAR_ vs TYPE_ ratios) added as a feature stream.')

# Summary table
heading('Summary — All 5 FS Experiment Runs', level=3)
tbl_sum = doc.add_table(rows=1, cols=7)
tbl_sum.style = 'Table Grid'
tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl_sum, ['Run', 'What Changed', 'Acc', 'F1', 'AUC', 'MCC', 'FPR'])
for row in [
    ('Run 1', 'Baseline — 4 rounds (32Q)',              '0.6440', '0.6555', '0.6939', '0.2884', '0.387'),
    ('Run 2', 'Retrained GAT — 8 rounds (64Q)',         '0.6724', '0.7129', '0.7064', '0.3580', '0.467'),
    ('Run 3', '8 rounds + full pipeline retrain',       '0.6514', '0.7273', '0.6822', '0.3609', '0.625'),
    ('Run 4', 'Retrained hybrid MLP',                   '0.6712', '0.7246', '0.7160', '0.3694', '0.521'),
    ('Run 5', '303-dim hybrid',                         '0.6761', '0.7277', '0.7131', '0.3784', '0.511'),
    ('FINAL', 'Supervisor improvements + XGB ensemble', '0.8195', '0.8274', '0.8900', '0.6412', '0.223'),
]:
    table_row(tbl_sum, row)
doc.add_paragraph()
body(
    'The quantum pipeline alone plateaued at ~67-68% accuracy (Run 5). The final jump to 82% '
    'came from adding enriched features, rule-based patterns, and an optimised XGBoost ensemble '
    '— incorporating the supervisor\'s guidance at every step.'
)

heading('5.6 Experiment 5 — VQC vs Classical Equivalent (Stage 6)', level=2)
body('Goal: Compare the VQC output to a classical MLP with the same number of trainable parameters (60).')
bullet('Classical MLP (60 params): accuracy ~64–65% on FS, ~82% on BO')
bullet('VQC (60 params): accuracy ~66% on FS, ~84% on BO')
bullet('VQC consistently outperformed the parameter-matched classical baseline by 1–2%')
body('The VQC captures non-linear quantum correlations between features — entanglement via CNOT-ring creates feature interactions that the classical MLP cannot replicate with the same parameter budget.')

heading('5.7 Experiment 6 — Data Re-uploading: 2 rounds vs 8 rounds', level=2)
body('Goal: Test whether increasing re-uploading rounds (and thus input feature count) improves accuracy.')
tbl13 = doc.add_table(rows=1, cols=4)
tbl13.style = 'Table Grid'
tbl13.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl13, ['Dataset', 'Rounds', 'Features Input', 'VQC Accuracy'])
for row in [
    ('BO',  '2',  '16', '~82%'),
    ('BO',  '8',  '64', '~84%'),
    ('FS',  '2',  '16', '~64%'),
    ('FS',  '8',  '64', '66.3%'),
    ('UAF', '2',  '16', '~80%'),
]:
    table_row(tbl13, row)
doc.add_paragraph()
body('Observation: More re-uploading rounds consistently improve VQC accuracy. The 8-round configuration (64 features) was selected as the final setting for BO and FS, following supervisor guidance.')

heading('5.8 Experiment 7 — Hybrid Fusion Ablation (Stage 7+8)', level=2)
body('Goal: Measure the contribution of each feature stream to final accuracy.')
tbl14 = doc.add_table(rows=1, cols=4)
tbl14.style = 'Table Grid'
tbl14.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl14, ['Feature Combination', 'BO Acc', 'FS Acc', 'UAF Acc'])
for row in [
    ('Classical only (256-dim)',                 '~89%', '~68%', '~88%'),
    ('Classical + VQC (288/260/311-dim)',        '~90%', '~70%', '~88%'),
    ('Classical + VQC + Rule + Enriched (full)','90.5%','82.0%','91.8%'),
]:
    table_row(tbl14, row)
doc.add_paragraph()
body('Each feature stream contributes incrementally. Rule-based features (format string patterns, pointer analysis) provided the largest single boost for FS due to the pattern-matching nature of format string vulnerabilities.')

heading('5.9 Experiment 8 — Ensemble Classifier Comparison', level=2)
body('Goal: Compare individual classifiers against the blended ensemble.')
tbl15 = doc.add_table(rows=1, cols=4)
tbl15.style = 'Table Grid'
tbl15.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl15, ['Classifier', 'BO Val Acc', 'FS Val Acc', 'UAF Val Acc'])
for row in [
    ('XGBoost alone',              '~87%', '~66%', '~85%'),
    ('GradientBoosting alone',     '~86%', '~65%', '~84%'),
    ('RandomForest alone',         '~85%', '~67%', '~83%'),
    ('Blended Ensemble (optimised)','~90%','~68%', '~89%'),
]:
    table_row(tbl15, row)
doc.add_paragraph()
body('The blended ensemble with grid-searched weights and threshold consistently outperforms any single classifier by 1–3%.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. Final Results')

heading('6.1 Test Set Performance', level=2)
tbl16 = doc.add_table(rows=1, cols=8)
tbl16.style = 'Table Grid'
tbl16.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl16, ['Dataset', 'Accuracy', 'Precision', 'Recall', 'F1', 'AUC', 'MCC', 'Threshold'])
for row in [
    ('BO',  '90.5%', '85.1%', '98.1%', '0.912', '0.972', '0.820', '0.344'),
    ('FS',  '82.0%', '79.6%', '86.2%', '0.827', '0.890', '0.641', '0.420'),
    ('UAF', '91.8%', '88.4%', '95.3%', '0.917', '0.938', '0.838', '0.154'),
]:
    table_row(tbl16, row)
doc.add_paragraph()

heading('6.2 Comparison with QEGVD Paper Baseline', level=2)
tbl17 = doc.add_table(rows=1, cols=4)
tbl17.style = 'Table Grid'
tbl17.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl17, ['Dataset', 'QEGVD Paper Baseline', 'SecuraQpp v2', 'Improvement'])
for row in [
    ('BO',  '78.6%', '90.5%', '+11.9%'),
    ('FS',  '61.2%', '82.0%', '+20.8%'),
    ('UAF', '74.3%', '91.8%', '+17.5%'),
]:
    table_row(tbl17, row)
doc.add_paragraph()

heading('6.3 Confusion Matrices', level=2)
body('Confusion matrix results for the final test set:')
tbl18 = doc.add_table(rows=1, cols=6)
tbl18.style = 'Table Grid'
tbl18.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl18, ['Dataset', 'TP', 'FP', 'TN', 'FN', 'FPR'])
for row in [
    ('BO',  '836', '146', '715', '16',  '16.9%'),
    ('FS',  '350', '90',  '313', '56',  '22.3%'),
    ('UAF', '61',  '8',   '62',  '3',   '11.4%'),
]:
    table_row(tbl18, row)
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. CHARTS
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. Graphs & Visualizations')

heading('7.1 Pipeline Progression — Accuracy per Stage', level=2)
add_img(CHARTS/'pipeline_progression.png', width=5.5,
        caption='Figure 1: Accuracy improvement from Stage 3 (GAT) through Stage 8 (Hybrid Fusion) for all three vulnerability types')

doc.add_paragraph()
heading('7.2 Final Metrics Comparison', level=2)
add_img(CHARTS/'metrics_comparison.png', width=5.5,
        caption='Figure 2: Side-by-side comparison of Accuracy, F1, AUC, and MCC across BO, FS, and UAF')

doc.add_paragraph()
heading('7.3 Training Loss Curves', level=2)
add_img(CHARTS/'training_curves.png', width=5.5,
        caption='Figure 3: Training and validation loss curves for the hybrid fusion MLP across all three pipelines')

doc.add_paragraph()
heading('7.4 Confusion Matrices', level=2)
add_img(CHARTS/'confusion_matrices.png', width=5.5,
        caption='Figure 4: Confusion matrices for BO, FS, and UAF on the held-out test set')

doc.add_paragraph()
heading('7.5 Metrics Summary Table', level=2)
add_img(CHARTS/'metrics_table.png', width=5.5,
        caption='Figure 5: Complete metrics table — Accuracy, Precision, Recall, F1, MCC, AUC per dataset')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. Web Application Architecture')

body(
    'SecuraQpp v2 is deployed as a three-tier web application. The ML pipeline runs as a '
    'separate FastAPI microservice, decoupled from the authentication layer.'
)

tbl19 = doc.add_table(rows=1, cols=3)
tbl19.style = 'Table Grid'
tbl19.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(tbl19, ['Layer', 'Technology', 'Responsibilities'])
for row in [
    ('Frontend',          'React + Vite',          'UI, force-directed graph viz (D3), scan history, PDF reports'),
    ('Scanning Backend',  'FastAPI (Python)',       'ML inference, quantum pipeline, vulnerability classification'),
    ('Auth Backend',      'Node.js + Express',      'Login, 2FA OTP (Nodemailer), JWT session management'),
    ('Database',          'SQLite (sql.js)',         'Users, roles, audit log, scan history'),
    ('AI Patch Engine',   'Claude API (Anthropic)', 'CWE-mapped vulnerability explanations and fix suggestions'),
]:
    table_row(tbl19, row)
doc.add_paragraph()

body('Key features:')
bullet('Role-based access control: Admin portal (purple) and Analyst portal (gold)')
bullet('Email OTP two-factor authentication for all logins')
bullet('Force-directed SVG graph visualization of code structure in the browser')
bullet('Real-time vulnerability scanning with confidence scores and severity levels')
bullet('PDF report export with full metric breakdown per scan')
bullet('AI patch suggestions mapped to CWE categories via Claude API')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading('9. Conclusion')
body(
    'SecuraQpp v2 successfully demonstrates that quantum-enhanced machine learning can '
    'outperform classical-only approaches for software vulnerability detection. The key '
    'contributions are:'
)
bullet('7-view multi-graph construction that captures structural, flow, and semantic code properties simultaneously')
bullet('QAFA composite scoring that selects the most vulnerability-relevant features before quantum encoding')
bullet('Data re-uploading VQC allowing 64 input features into a 4-qubit circuit through 8 encoding rounds')
bullet('Hybrid ensemble fusion combining quantum, classical, enriched, and rule-based signals')
bullet('End-to-end deployed system with 2FA auth, role-based access, and AI patch suggestions')
body(
    'Final test accuracies of 90.5% (BO), 91.8% (UAF), and 82.0% (FS) represent improvements '
    'of +11.9%, +17.5%, and +20.8% respectively over the QEGVD paper baseline. Future work '
    'includes deployment on real quantum hardware (IBM/IonQ) and extending to inter-procedural '
    'taint analysis for improved FS detection.'
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f'Saved: {OUT}')
